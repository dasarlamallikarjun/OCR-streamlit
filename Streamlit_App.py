import streamlit as st
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from io import BytesIO
from docx import Document

class PDF_OCR:
    def __init__(self, pdf_path, lang="eng"):
        self.pdf_path = pdf_path
        self.lang = lang

    def process_pdf(self):
        # Open the PDF file
        try:
            doc = fitz.open(self.pdf_path)
        except Exception as e:
            st.error(f"Error opening PDF: {e}")
            return None
        
        # Create a new Word document
        word_doc = Document()
        pages_processed = 0
        text_found = 0
        
        # Process each page in the PDF
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")
            
            if text.strip():  # Text was found
                text_found += 1
                word_doc.add_paragraph(text)
            else:  # If no text, perform OCR using Tesseract
                pix = page.get_pixmap()
                img = Image.open(BytesIO(pix.tobytes("png")))
                ocr_text = pytesseract.image_to_string(img, lang=self.lang)
                
                if ocr_text.strip():
                    word_doc.add_paragraph(ocr_text)

            pages_processed += 1
        
        # Save the Word document
        output_path = "extracted_text.docx"
        word_doc.save(output_path)

        # Show summary message
        st.success(f"Processing complete: {pages_processed} pages processed, {text_found} pages had extractable text.")
        
        return output_path

st.title("PDF to Text OCR Application")

# Language selection
language_options = {
    "English": "eng",
    "Telugu": "tel",
    "Hindi": "hin",
    "Spanish": "spa",  # You can add more languages as needed
    "French": "fra"
}

selected_lang = st.selectbox("Select Language for OCR", list(language_options.keys()))

# File upload
uploaded_file = st.file_uploader("Upload a PDF file", type=["pdf"])

if uploaded_file is not None:
    # Save the uploaded PDF temporarily
    temp_pdf_path = "temp_uploaded_file.pdf"
    with open(temp_pdf_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    
    # Initialize and process PDF with OCR
    pdf_ocr = PDF_OCR(temp_pdf_path, lang=language_options[selected_lang])  # Use selected language
    output_path = pdf_ocr.process_pdf()

    if output_path:
        # Show download button for DOCX file
        with open(output_path, "rb") as file:
            st.download_button(
                label="Download Extracted Text as DOCX",
                data=file,
                file_name="extracted_text.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.error("Failed to process PDF.")
